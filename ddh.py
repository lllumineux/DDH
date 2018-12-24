import json
import sys
from PyQt5 import QtCore, QtGui
from PyQt5.QtWidgets import QInputDialog, QMessageBox, QWidget, QGridLayout, \
    QVBoxLayout, QApplication, QMainWindow, QDesktopWidget, QPushButton, \
    QSizePolicy, QLabel, QScrollArea, QLineEdit, QHBoxLayout
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt


# Главное окно
class ddh(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        # Создание главного окна
        x = (QDesktopWidget().screenGeometry().width() -
             QDesktopWidget().screenGeometry().width() // 2) // 2
        y = (QDesktopWidget().screenGeometry().height() -
             QDesktopWidget().screenGeometry().height() // 2) // 2
        self.setGeometry(x, y, QDesktopWidget().screenGeometry().width() / 2,
                         QDesktopWidget().screenGeometry().height() / 2)
        self.setObjectName('main_window')
        self.setWindowTitle('DDH')
        self.setStyleSheet('background: #fafafa;')

        # Создание виджета главного окна
        self.main_window_widget = QWidget(self)
        self.main_window_widget.setObjectName('main_window_widget')
        self.main_window_widget.setStyleSheet('background: #fafafa;')

        # Создание сетки главного окна
        self.main_window_grid = QGridLayout(self.main_window_widget)
        self.main_window_grid.setObjectName('main_window_grid')

        # Создание виджета для кнопок
        self.btn_widget = QWidget(self.main_window_widget)
        self.btn_widget.setObjectName('btn_widget')
        self.btn_widget.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        self.btn_widget.setMinimumSize(QtCore.QSize(300, 270))
        self.btn_widget.setMaximumSize(QtCore.QSize(300, 270))

        # Создание блока кнопок
        self.btn_block = QVBoxLayout(self.btn_widget)
        self.btn_block.setObjectName('btn_block')
        self.btn_block.setContentsMargins(50, 25, 50, 25)
        self.btn_block.setSpacing(35)

        # Создание кнопок для редактирования json файла
        btn_create_info = (('add', 'добавить', self.add_func),
                           ('change', 'изменить', self.change_func),
                           ('del', 'удалить', self.del_func),
                           ('exp', 'экспортировать', self.exp_func))

        for obj_name, name, btn_func in btn_create_info:
            self.btn = QPushButton(self.btn_widget)
            self.btn.setObjectName('btn_' + obj_name)
            self.btn.setText(name.upper())
            self.btn.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
            self.btn.setMinimumSize(QtCore.QSize(200, 50))
            self.btn_block.addWidget(self.btn)
            self.btn.clicked.connect(btn_func)
            self.btn.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
            self.btn.setStyleSheet('QPushButton {background: #212121; color: #fafafa;'
                                   'font-family: Arial Black; font-size: 10px;'
                                   'border: none;}'
                                   'QPushButton:hover {background: #383838;}')

        # Добавление сеток друг в друга
        self.main_window_grid.addWidget(self.btn_widget, 1, 1, 1, 1)
        self.setCentralWidget(self.main_window_widget)

    # Добавление функции в json файл
    def add_func(self):
        ans1, ok_pressed1 = QInputDialog.getText(
            self, 'Добавить функцию', 'Введите название функции')
        ans2, ok_pressed2, wrong_value = '', False, False

        try:
            ans2, ok_pressed2 = QInputDialog.getText(
                self, 'Добавить функцию', 'Введите количество аргументов')

            if ans2.isalpha() or int(ans2) < 0:
                QMessageBox.information(self, 'Ошибка', 'Неверное значение!')
                wrong_value = True

        except Exception:
            pass

        if ok_pressed1 and ok_pressed2 and not wrong_value and ans2 != '':
            self.dialog = func_window(ans1, int(ans2))
            self.dialog.show()

    # Изменение существующей функции из json файла
    def change_func(self):
        ans, ok_pressed = QInputDialog.getText(
            self, 'Изменить функцию', 'Введите название функции')

        try:
            with open('base.json', 'r', encoding='utf-8') as file:
                base = json.load(file)

                if not ans in base:
                    QMessageBox.information(self, 'Ошибка', 'Функция не найдена!')
                    base = {}
        except Exception:
            base = {}

        if ok_pressed and ans in base:
            self.dialog = func_window(ans)
            self.dialog.show()

    # Удаление функции из json файла
    def del_func(self):
        ans, ok_pressed = QInputDialog(self).getText(
            self, 'Удалить функцию', 'Введите название функции')

        try:
            with open('base.json', 'r', encoding='utf-8') as file:
                base = json.load(file)
                if ok_pressed:
                    del base[ans]

            with open('base.json', 'w+', encoding='utf-8') as file:
                json.dump(base, file)

            del_result_text = ('Удалено', 'Функция успешно удалена')

        except Exception:
            del_result_text = ('Ошибка', 'Функция не найдена')

        if ans != '':
            QMessageBox.information(self, *del_result_text)

    # Выбор формата для экспорта json файла
    def exp_func(self):
        exp_format, ok_pressed = QInputDialog.getItem(
            self, 'Экспортировать', 'Выберите формат экспорта',
            ('html', 'docx'), 0, False)

        if ok_pressed:
            try:
                with open('base.json', 'r', encoding='utf-8') as file:
                    base = json.load(file)
            except Exception:
                base = {}

            if exp_format == 'html':
                html_file = '''<!DOCTYPE html><html><head><meta http-equiv="content-type" content="text/html"><meta name="viewport" content="width=device-width, initial-scale=1, shrink-to-fit=no"><meta charset="utf-8"><title>DDH</title><style type="text/css">@import url('https://fonts.googleapis.com/css?family=montserrat:300,400,500,600');html,body{font-family:'montserrat',sans-serif;overflow-x:hidden;width:100%;padding:0;margin:0;-webkit-transition:all 0.3s ease;-moz-transition:all 0.3s ease;-o-transition:all 0.3s ease;transition:all 0.3s ease}header,main,footer,section,nav,ul,li{display:block;padding:0;margin:0;z-index:1}h1,h2,h3,h4,h5,h6,p,i,b,a,textarea,input,button{text-decoration:none;outline:0;padding:0;margin:0}a:hover,textarea:hover,input:hover,button:hover{text-decoration:none}textarea:focus,input:focus,button:focus{outline:0}.close-search-overlay{position:fixed;top:0;right:0;left:0;bottom:0;width:100vw;height:100vh}.search-overlay{display:flex;flex-direction:column;justify-content:center;align-items:center;overflow-y:hidden;-webkit-transition:all 0.3s ease;-moz-transition:all 0.3s ease;-o-transition:all 0.3s ease;transition:all 0.3s ease;background:#000;position:fixed;top:0;right:0;left:0;bottom:0;width:100vw;height:100vh;color:#fff;z-index:2;opacity:0.85}.search-error-p{-webkit-transition:all 0.3s ease;-moz-transition:all 0.3s ease;-o-transition:all 0.3s ease;transition:all 0.3s ease;opacity:0;z-index:-1}.search-overlay__element{display:flex;flex-direction:column;justify-content:center;align-items:center;margin:0 0 77.5px 0}.search-overlay input{border:none;background:transparent;border-bottom:2px solid #000;color:#fff;font-size:27.5px;font-weight:300}.search-overlay p{color:#fff;font-size:50px;font-weight:300;margin:0 0 50px 0}.search{display:flex;justify-content:flex-end;align-items:center;margin:0 70px 0 0;height:100px;opacity:0}.open-search-overlay{-webkit-transition:all 0.5s ease;-moz-transition:all 0.5s ease;-o-transition:all 0.5s ease;transition:all 0.5s ease}.open-search-overlay:hover{opacity:0.5}.open-search-overlay img{width:40px}.func{display:flex;flex-direction:column;justify-content:center;align-items:center;margin:35px}.func-name h2{text-align:center;font-weight:400;font-size:35px}.func h3{font-weight:500;font-size:22.5px}.func h4{color:#5a5a5a;font-weight:600;text-align:center;font-size:17.5px}.func h5{color:#5a5a5a;font-weight:600;font-size:17.5px}.func-p{font-weight:300;font-size:17.5px}.arg-p{color:#5a5a5a;font-weight:400;font-size:15px;margin:0 0 0 15px}.description,.syntax,.return-value{width:1000px;margin:0 0 15px 0}.args{width:1000px;margin:0 0 5px 0}@media (max-width:1100px){.description,.syntax,.args,.return-value{width:90%}}.func__element{margin:5px 0 5px 15px}.arg{margin:15px 50px}@media (max-width:750px){.arg{margin:15px 0}}</style></head><body><nav class="search"><a href="#" class="open-search-overlay"><img src="search.png"></a></nav><main class="main"><div class="main__element"><div class="func-name"><h2></h2></div><div class="func"><div class="description"></div><div class="syntax"></div><div class="args"></div><div class="return-value"></div></div></div></main><div class="search-overlay"><div class="search-overlay__element"><p class="search-error-p">Функция не найдена!</p><input id="search-form" type="text" name="func-search" placeholder="Введите имя_функции()"></div></div><script src="https://code.jquery.com/jquery-3.3.1.slim.min.js" integrity="sha384-q8i/X+965DzO0rT7abK41JStQIAqVgRVzpbzo5smXKp4YfRvH+8abtTE1Pi6jizo" crossorigin="anonymous"></script><script type="text/javascript">base=place-to-paste-json-file;contentonPage=false;$(document).ready(()=>{$('#search-form')[0].focus();});$(document).keyup((e)=>{$('.search-error-p').css('opacity', '0');$('.search-error-p').css('z-index', '-1');if(e.which==13){for(func_name in base){if(func_name==$('#search-form')[0].value){$('.search-overlay').css('opacity', '0');$('.search-overlay').css('z-index', '-1');func_info(base[func_name], func_name);$('.search').css('opacity', '1');contentonPage = true;break}else{$('.search-error-p').css('opacity', '1');$('.search-error-p').css('z-index', '2');};};$('#search-form')[0].value = '';};});$('.open-search-overlay').click(()=>{$('.search-overlay').css('opacity', '0.85');$('.search-overlay').css('z-index', '2');$('.search-error-p').css('opacity', '0');$('.search-error-p').css('z-index', '-1');$('#search-form')[0].focus();});$('.search-overlay').click(()=>{if (contentonPage == true){$('.search-overlay').css('opacity', '0');$('.search-overlay').css('z-index', '-1');};});function func_info(func_info, func_name){if(func_name!=''){$('.func-name')[0].innerHTML='<h2>'+func_name+'</h2>';};if(func_info['description']!=''){$('.description')[0].innerHTML='<h3>Описание</h3><p class="func__element func-p">'+func_info['description']+'</p>';};if(func_info['syntax']!=''){$('.syntax')[0].innerHTML='<h3>Синтаксис</h3><p class="func__element func-p">'+func_info['syntax']+'</p>';};if(func_info['args']!={}){$('.args')[0].innerHTML='<h3>Аргументы</h3><div class="func__element func_args"></div>';};if (func_info['return_value']!=''){$('.return-value')[0].innerHTML='<h3>Возвращаемое значение</h3><p class="func__element func-p func_return-value">'+func_info['return_value']+'</p>';};args_info='';for (arg in func_info['args']){arg_info='<div class="arg">';arg_info+='<div class="arg__name"><h4>'+arg+'</h4></div>';arg_info+='<div class="arg__description"><h5>Описание</h5><p class="arg-p">'+func_info['args'][arg]['description']+'</p></div>';arg_info+='<div class="arg__type"><h5>Тип</h5><p class="arg-p">'+func_info['args'][arg]['type']+'</p></div>';arg_info+='<div class="arg__default-value"><h5>Значение по умолчанию</h5><p class="arg-p">'+func_info['args'][arg]['default_value']+'</p></div>';arg_info+='<div class="arg__addition"><h5>Примечание</h5><p class="arg-p">'+func_info['args'][arg]['addition']+'</p></div></div>';args_info+=arg_info;};$('.func_args')[0].innerHTML = args_info;};</script></body></html>'''\
                    .replace('place-to-paste-json-file', str(base))
                # Содержимое итогового html файла

                with open('ddh.html', 'w+', encoding='utf-8') as file:
                    file.write(html_file)

            elif exp_format == 'docx':
                # Создание документа
                doc = Document()

                # Стили документа
                styles = doc.styles

                # Изменение шрифта для заглавия
                style = styles.add_style('Big_size', WD_STYLE_TYPE.CHARACTER)
                big = doc.styles['Big_size']
                big_font = big.font
                big_font.name = 'Arial'
                big_font.size = Pt(24)

                # Изменение шрифта постепенно
                normal = doc.styles['Normal']
                font = normal.font
                font.name = 'Arial'
                font.size = Pt(16)
                style = styles.add_style('Citation', WD_STYLE_TYPE.PARAGRAPH)
                small = doc.styles['Citation']
                small_font = small.font
                small_font.name = 'Arial'
                small_font.size = Pt(14)
                style1 = styles.add_style('Smallest_small', WD_STYLE_TYPE.PARAGRAPH)
                smallest = doc.styles['Smallest_small']
                smallest_font = smallest.font
                smallest_font.name = 'Arial'
                smallest_font.size = Pt(12)

                # Добавление заглавия
                run = doc.add_paragraph()
                paragraph_format = run.paragraph_format
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.add_run('Справка по функциям', style='Big_size').bold = True

                # Списки функций, их аргументов и других элементов
                functions = []
                for i in base:
                    functions.append(i)
                elements = []
                for i in functions:
                    for j in base[i]:
                        elements.append(base[i][j])
                args = []
                for i in functions:
                    i = base[i]['args'].keys()
                    a = []
                    for j in i:
                        a.append(j)
                        if len(a) != 0:
                            args.append(a)
                elements_args = []
                for i in functions:
                    for j in base[i]['args']:
                        for g in base[i]['args'][j]:
                            elements_args.append(base[i]['args'][j][g])

                # Написание самого документа
                for k in range(len(functions)):
                    p = doc.add_paragraph()
                    p.add_run(functions[k]).bold = True
                    for key in base[functions[k]]:
                        if key == 'description':
                            p1 = doc.add_paragraph('Описание функции:',
                                                   style='Citation')
                        if key == 'syntax':
                            p1 = doc.add_paragraph('Синтакс функции:',
                                                   style='Citation')
                        if key == 'return_value':
                            p1 = doc.add_paragraph(
                                'Возвращаемое значение функции:',
                                style='Citation')
                        if key == 'args' and len(args) != 0:
                            p1 = doc.add_paragraph('Аргументы функции:',
                                                   style='Citation')
                        paragraph_format1 = p1.paragraph_format
                        for value1 in range(len(elements)):
                            if key == 'description' and value1 == 0:
                                p2 = doc.add_paragraph(elements[value1],
                                                       style='Citation')
                            if key == 'syntax' and value1 == 1:
                                p2 = doc.add_paragraph(elements[value1],
                                                       style='Citation')
                            if key == 'return_value' and value1 == 3:
                                p2 = doc.add_paragraph(elements[value1],
                                                       style='Citation')
                        paragraph_format1.first_line_indent = Inches(0.25)
                        paragraph_format2 = p2.paragraph_format
                        paragraph_format2.first_line_indent = Inches(0.50)
                        if key == 'args' and len(args) != 0:
                            for l in range(len(args[k])):
                                p3 = doc.add_paragraph(args[k][l] + ':',
                                                       style='Citation')
                                paragraph_format3 = p3.paragraph_format
                                paragraph_format3.first_line_indent = Inches(
                                    0.50)
                                for f in base[functions[k]]['args'][
                                    args[k][l]]:
                                    for value2 in range(len(elements_args)):
                                        if f == 'description' and value2 == 0:
                                            p4 = doc.add_paragraph(
                                                'Описане аргумента:',
                                                style='Smallest_small')
                                            p5 = doc.add_paragraph(
                                                elements_args[value2],
                                                style='Smallest_small')
                                        if f == 'type' and value2 == 1:
                                            p4 = doc.add_paragraph(
                                                'Тип аргумента:',
                                                style='Smallest_small')
                                            p5 = doc.add_paragraph(
                                                elements_args[value2],
                                                style='Smallest_small')
                                        if f == 'default_value' and value2 == 2:
                                            p4 = doc.add_paragraph(
                                                'Начальное значение аргумента:',
                                                style='Smallest_small')
                                            p5 = doc.add_paragraph(
                                                elements_args[value2],
                                                style='Smallest_small')
                                        if f == 'addition' and value2 == 3:
                                            p4 = doc.add_paragraph(
                                                'Примичание:',
                                                style='Smallest_small')
                                            p5 = doc.add_paragraph(
                                                elements_args[value2],
                                                style='Smallest_small')
                                        paragraph_format5 = p5.paragraph_format
                                        paragraph_format5.first_line_indent = Inches(
                                            1.0)
                                        paragraph_format4 = p4.paragraph_format
                                        paragraph_format4.first_line_indent = Inches(
                                            0.75)
                                elements_args = elements_args[4:]
                    elements = elements[4:]
                    p6 = doc.add_paragraph()
                doc.add_page_break()
                doc.save('ddh.docx')


# Окно редактирования json файла
class func_window(QMainWindow):
    def __init__(self, func_name_str, arg_amount=-1):
        super().__init__()
        self.arg_amount, self.func_name_str, self.args_info_form = \
            arg_amount, func_name_str, []
        self.is_change_func = True if arg_amount == -1 else False
        self.setupUi()

    def setupUi(self):
        # Параметры окна
        self.setObjectName('func_window')
        x = (QDesktopWidget().screenGeometry().width() -
             QDesktopWidget().screenGeometry().width() // 2) // 2
        y = (QDesktopWidget().screenGeometry().height() -
             QDesktopWidget().screenGeometry().height() // 2) // 2
        self.setGeometry(x, y, QDesktopWidget().screenGeometry().width() / 2,
                         QDesktopWidget().screenGeometry().height() / 2)
        self.setWindowTitle(self.func_name_str)
        self.setStyleSheet('background: #fafafa;')

        # Виджет окна
        self.func_window_grid = QWidget(self)
        self.func_window_grid.setObjectName('func_window_grid')

        # Основная сетка окна
        self.vertical_layout_1 = QVBoxLayout(self.func_window_grid)
        self.vertical_layout_1.setObjectName('vertical_layout_1')

        # Скролл-область
        self.scroll_area = QScrollArea(self.func_window_grid)
        self.scroll_area.setObjectName('scroll_area')
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area_widget = QWidget(self)
        self.scroll_area_widget.setObjectName('scroll_area_widget')

        # Блок параметров
        self.vertical_layout_2 = QVBoxLayout(self.scroll_area_widget)
        self.vertical_layout_2.setObjectName('vertical_layout_2')

        # Блок описания и добавление его в скролл-область
        self.func_description_layout_2 = QVBoxLayout()
        self.func_description_layout_2.setObjectName('func_description_layout_2')
        self.func_description_label_2 = QLabel(self.scroll_area_widget)

        # Название блока описания
        self.func_description_label_2.setObjectName('func_description_label_2')
        self.func_description_label_2.setText('Описание')
        self.func_description_layout_2.addWidget(self.func_description_label_2)

        # Форма редактирования блока описания
        self.func_description_edit_2 = QLineEdit(self.scroll_area_widget)
        self.func_description_edit_2.setObjectName('func_description_edit_2')
        self.func_description_layout_2.addWidget(self.func_description_edit_2)

        # Добавление блока описания в блок параметров
        self.vertical_layout_2.addLayout(self.func_description_layout_2)

        # Блок синтаксиса и добавление его в скролл-область
        self.func_syntax_layout_2 = QVBoxLayout()
        self.func_syntax_layout_2.setObjectName('func_syntax_layout_2')
        self.func_syntax_label_2 = QLabel(self.scroll_area_widget)

        # Название блока синтаксиса
        self.func_syntax_label_2.setObjectName('func_syntax_label_2')
        self.func_syntax_label_2.setText('Синтаксис')
        self.func_syntax_layout_2.addWidget(self.func_syntax_label_2)

        # Форма редактирования блока синтаксиса
        self.func_syntax_edit_2 = QLineEdit(self.scroll_area_widget)
        self.func_syntax_edit_2.setObjectName('func_syntax_edit_2')
        self.func_syntax_layout_2.addWidget(self.func_syntax_edit_2)

        # Добавление блока синтаксиса в блок параметров
        self.vertical_layout_2.addLayout(self.func_syntax_layout_2)

        # Блок возвращаемого значения и добавление его в скролл-область
        self.func_return_value_layout_2 = QVBoxLayout()
        self.func_return_value_layout_2.setObjectName('func_return_value_layout_2')
        self.func_return_value_label_2 = QLabel(self.scroll_area_widget)

        # Название возвращаемого значения
        self.func_return_value_label_2.setObjectName('func_return_value_label_2')
        self.func_return_value_label_2.setText('Возвращаемое значение')
        self.func_return_value_layout_2.addWidget(self.func_return_value_label_2)

        # Форма редактирования возвращаемого значения
        self.func_return_value_edit_2 = QLineEdit(self.scroll_area_widget)
        self.func_return_value_edit_2.setObjectName('func_return_value_edit_2')
        self.func_return_value_layout_2.addWidget(self.func_return_value_edit_2)

        # Добавление блока возвращаемого значения в блок параметров
        self.vertical_layout_2.addLayout(self.func_return_value_layout_2)

        # Добавление скролл-области в основную сетку окна
        self.scroll_area.setWidget(self.scroll_area_widget)
        self.vertical_layout_1.addWidget(self.scroll_area)

        # Блок навигации
        self.func_window_nav = QHBoxLayout()
        self.func_window_nav.setObjectName('func_window_nav')

        # Кнопка 'сохравнить и выйти'
        self.save_btn = QPushButton(self.func_window_grid)
        self.save_btn.setObjectName('save_btn')
        self.save_btn.setText('Сохравнить и выйти')
        self.save_btn.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        self.save_btn.setMinimumSize(QtCore.QSize(200, 30))
        self.save_btn.setMaximumSize(QtCore.QSize(200, 30))
        self.save_btn.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
        self.save_btn.setStyleSheet('QPushButton {background: #212121;'
                                             ' color: #fafafa; font-family: Arial Black; '
                                             'font-size: 10px; border: none;} '
                                             'QPushButton:hover {background: #383838;}')
        self.save_btn.clicked.connect(self.save)

        # Добавление кнопки 'сохравнить и выйти' в блок навигации
        self.func_window_nav.addWidget(self.save_btn)

        # Добавление блока навигации в основную сетку окна
        self.vertical_layout_1.addLayout(self.func_window_nav)

        # Проверка колличества аргументов
        if self.arg_amount > 0:
            # Блок аргументов
            self.args_layout = QVBoxLayout()
            self.args_layout.setObjectName('args_layout')

            # Название блока аргументов
            self.label_args = QLabel(self.scroll_area_widget)
            self.label_args.setObjectName('label_args')
            self.label_args.setText('Аргументы')
            self.args_layout.addWidget(self.label_args)

            # Добавление блока аргументов в блок параметров
            self.vertical_layout_2.addLayout(self.args_layout)

        # Назначение основного виджета окна
        self.setCentralWidget(self.func_window_grid)

        if self.is_change_func:
            try:
                with open('base.json', 'r', encoding='utf-8') as file:
                    base = json.load(file)
                self.arg_amount = len(base[self.func_name_str]['args'])
            except Exception:
                QMessageBox.information(self, 'Ошибка', 'Функция не найдена!')

        self.n = 0
        for i in range(self.arg_amount):
            self.add_arg()

        if self.is_change_func:
            self.fill_window()

# Добавление одного аргумента
    def add_arg(self):
        self.n += 1
        self.a = 0
        self.grid_args = QVBoxLayout()
        self.label_name_arg = QLabel('аргумент #' + str(self.n), self.scroll_area_widget)
        self.grid_args.addWidget(self.label_name_arg)
        self.vertical_layout_2.addLayout(self.grid_args)
        self.label_name_arg.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        self.vertical_layout_2.addLayout(self.grid_args)
        arg_info = (('Название аргумента', 'arg_name_', 'name_arg_lineEdit_'),
                    ('Описаниее аргумента', 'label_description_arg_', 'description_arg_lineEdit_'),
                    ('Тип аргумента', 'label_type_arg_', 'type_arg_lineEdit_'),
                    ('Начальное значение аргумента', 'label_default_value_', 'default_value_lineEdit_'),
                    ('Примечание', 'label_additions', 'additions_lineEdit_'))
        self.arg_info_form = []
        for name, label, lineEdit in arg_info:
            self.grid_arg = QVBoxLayout()
            self.arg_name = QLabel(name, self.scroll_area_widget)
            self.grid_arg.addWidget(self.arg_name)
            self.arg_lineEdit = QLineEdit()
            self.grid_arg.addWidget(self.arg_lineEdit)
            self.arg_name.setObjectName(label + str(self.a))
            self.arg_lineEdit.setObjectName(lineEdit + str(self.a))
            self.vertical_layout_2.addLayout(self.grid_arg)
            self.a += 1
            self.arg_info_form.append((self.arg_lineEdit))

        self.args_info_form.append(self.arg_info_form)

# Сохранение введённых данных
    def save(self):
        try:
            with open('base.json', 'r', encoding='utf-8') as file:
                base = json.load(file)
        except Exception:
            base = {}

        func = {
            self.func_name_str: {
                'description': self.func_description_edit_2.text(),
                'syntax': self.func_syntax_edit_2.text(),
                'args': {
                     self.args_info_form[i][0].text(): {
                        'description': self.args_info_form[i][1].text(),
                        'type': self.args_info_form[i][2].text(),
                        'default_value': self.args_info_form[i][3].text(),
                        'addition': self.args_info_form[i][4].text()
                     }
                     for i in range(len(self.args_info_form))
                },
                'return_value': self.func_return_value_edit_2.text()
            }
        }

        base[self.func_name_str] = func[self.func_name_str]

        with open('base.json', 'w+', encoding='utf-8') as file:
            json.dump(base, file)

        self.close()

# Заполнение параметров функции, после нажатия на кнопку изменить
    def fill_window(self):
        try:
            with open('base.json', 'r', encoding='utf-8') as file:
                base = json.load(file)

            self.func_description_edit_2.setText(base[self.func_name_str]['description'])
            self.func_syntax_edit_2.setText(base[self.func_name_str]['syntax'])
            self.func_return_value_edit_2.setText(base[self.func_name_str]['return_value'])
            i = 0
            for arg in base[self.func_name_str]['args']:
                self.args_info_form[i][0].setText(arg)
                self.args_info_form[i][1].setText(base[self.func_name_str]['args'][arg]['description'])
                self.args_info_form[i][2].setText(base[self.func_name_str]['args'][arg]['type'])
                self.args_info_form[i][3].setText(base[self.func_name_str]['args'][arg]['default_value'])
                self.args_info_form[i][4].setText(base[self.func_name_str]['args'][arg]['addition'])
                i += 1

        except Exception:
            QMessageBox.information(self, 'Функция не найдена!')


if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = ddh()
    ex.show()
    sys.exit(app.exec_())
